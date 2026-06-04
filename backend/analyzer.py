"""
analyzer.py
-----------
Core resume analysis engine.
- Text extraction from PDF / DOCX
- Field extraction (name, email, phone, links)
- Section detection (education, experience, projects, certifications)
- Skill detection & matching against job role requirements
- ATS score calculation with section-wise breakdown
- Language detection
- Formatting analysis
- Dictionary-based spelling checker
"""

import re
import os
from collections import OrderedDict

# ═══════════════════════════════════════════════════════════════
#  JOB ROLE → REQUIRED SKILLS MAPPING
# ═══════════════════════════════════════════════════════════════

JOB_ROLES = OrderedDict({
    "Data Analyst": [
        "Python", "SQL", "Excel", "Tableau", "Power BI", "Pandas",
        "NumPy", "Statistics", "Data Visualization", "R", "ETL",
        "Data Cleaning", "Matplotlib", "Seaborn"
    ],
    "Data Scientist": [
        "Python", "R", "SQL", "Machine Learning", "Deep Learning",
        "TensorFlow", "PyTorch", "Pandas", "NumPy", "Scikit-learn",
        "Statistics", "NLP", "Data Visualization", "Tableau", "Keras"
    ],
    "Business Analyst": [
        "SQL", "Excel", "Tableau", "Power BI", "JIRA", "Agile",
        "Requirement Gathering", "Data Analysis", "Stakeholder Management",
        "Business Intelligence", "Visio", "Communication"
    ],
    "Machine Learning Engineer": [
        "Python", "TensorFlow", "PyTorch", "Scikit-learn", "Keras",
        "Deep Learning", "Machine Learning", "NLP", "Computer Vision",
        "Docker", "MLOps", "AWS", "SQL", "Pandas", "NumPy"
    ],
    "AI Engineer": [
        "Python", "TensorFlow", "PyTorch", "Deep Learning", "NLP",
        "Computer Vision", "Machine Learning", "Keras", "OpenCV",
        "Transformers", "LLM", "Docker", "AWS", "MLOps"
    ],
    "Python Developer": [
        "Python", "Django", "Flask", "FastAPI", "SQL", "PostgreSQL",
        "REST API", "Git", "Docker", "Unit Testing", "OOP",
        "Celery", "Redis", "Linux"
    ],
    "Java Developer": [
        "Java", "Spring Boot", "Hibernate", "Maven", "SQL", "MySQL",
        "REST API", "Microservices", "JUnit", "Git", "Docker",
        "Kafka", "Jenkins", "OOP"
    ],
    "Software Developer": [
        "Python", "Java", "C++", "SQL", "Git", "REST API", "OOP",
        "Data Structures", "Algorithms", "Docker", "Linux",
        "Agile", "Unit Testing", "CI/CD"
    ],
    "Frontend Developer": [
        "HTML", "CSS", "JavaScript", "React", "TypeScript", "Redux",
        "Tailwind CSS", "Bootstrap", "Responsive Design", "Git",
        "Webpack", "Figma", "REST API", "Jest"
    ],
    "Backend Developer": [
        "Python", "Java", "Node.js", "SQL", "PostgreSQL", "MongoDB",
        "REST API", "Docker", "Git", "Redis", "Microservices",
        "Linux", "CI/CD", "AWS"
    ],
    "Full Stack Developer": [
        "HTML", "CSS", "JavaScript", "React", "Node.js", "Python",
        "SQL", "MongoDB", "REST API", "Git", "Docker", "TypeScript",
        "AWS", "CI/CD"
    ],
    "Web Developer": [
        "HTML", "CSS", "JavaScript", "React", "PHP", "MySQL",
        "Bootstrap", "WordPress", "Git", "Responsive Design",
        "jQuery", "REST API", "SEO", "Figma"
    ],
    "Android Developer": [
        "Java", "Kotlin", "Android SDK", "XML", "Firebase", "REST API",
        "SQLite", "Git", "MVVM", "Retrofit", "Jetpack Compose",
        "Material Design", "Gradle", "Unit Testing"
    ],
    "UI/UX Designer": [
        "Figma", "Adobe XD", "Sketch", "Wireframing", "Prototyping",
        "User Research", "Usability Testing", "Design Thinking",
        "Adobe Photoshop", "Adobe Illustrator", "HTML", "CSS",
        "Responsive Design", "Typography"
    ],
    "Cloud Engineer": [
        "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform",
        "Linux", "CI/CD", "Networking", "IAM", "CloudFormation",
        "Ansible", "Python", "Monitoring"
    ],
    "DevOps Engineer": [
        "Docker", "Kubernetes", "Jenkins", "CI/CD", "AWS", "Linux",
        "Terraform", "Ansible", "Git", "Python", "Bash",
        "Monitoring", "Prometheus", "Grafana"
    ],
    "Cyber Security Analyst": [
        "Network Security", "Penetration Testing", "SIEM", "Firewall",
        "Vulnerability Assessment", "Linux", "Python", "Wireshark",
        "Encryption", "IDS/IPS", "SOC", "Incident Response",
        "OWASP", "Ethical Hacking"
    ],
    "Database Administrator": [
        "SQL", "MySQL", "PostgreSQL", "Oracle", "MongoDB", "Database Design",
        "Backup Recovery", "Performance Tuning", "Linux", "Python",
        "Replication", "NoSQL", "Data Modeling", "Shell Scripting"
    ],
    "System Administrator": [
        "Linux", "Windows Server", "Active Directory", "Networking",
        "DNS", "DHCP", "Bash", "PowerShell", "VMware",
        "Backup Recovery", "Monitoring", "Security", "Firewall", "Docker"
    ],
    "Network Engineer": [
        "Cisco", "Routing", "Switching", "TCP/IP", "DNS", "DHCP",
        "Firewall", "VPN", "VLAN", "BGP", "OSPF", "Linux",
        "Network Security", "Wireshark"
    ],
    "QA Tester": [
        "Manual Testing", "Test Cases", "Bug Tracking", "JIRA",
        "Selenium", "Regression Testing", "API Testing", "SQL",
        "Agile", "Test Plan", "Smoke Testing", "UAT",
        "Performance Testing", "Postman"
    ],
    "Automation Tester": [
        "Selenium", "Java", "Python", "TestNG", "Cucumber", "CI/CD",
        "Jenkins", "REST API Testing", "Postman", "Git",
        "JUnit", "Maven", "Appium", "Performance Testing"
    ],
    "Digital Marketing Executive": [
        "SEO", "SEM", "Google Ads", "Facebook Ads", "Social Media Marketing",
        "Content Marketing", "Email Marketing", "Google Analytics",
        "WordPress", "Canva", "Copywriting", "Marketing Strategy"
    ],
    "Content Writer": [
        "Content Writing", "SEO Writing", "Blog Writing", "Copywriting",
        "WordPress", "Research", "Editing", "Proofreading",
        "Social Media", "Content Strategy", "Grammar", "Creative Writing"
    ],
    "Graphic Designer": [
        "Adobe Photoshop", "Adobe Illustrator", "Figma", "Canva",
        "InDesign", "Typography", "Branding", "Logo Design",
        "UI Design", "Color Theory", "Print Design", "After Effects"
    ],
    "Project Manager": [
        "Project Planning", "Agile", "Scrum", "JIRA", "Stakeholder Management",
        "Risk Management", "Budgeting", "MS Project", "Communication",
        "Leadership", "PMP", "Gantt Chart"
    ],
    "Product Manager": [
        "Product Strategy", "Roadmap", "Agile", "JIRA", "User Research",
        "A/B Testing", "Data Analysis", "Stakeholder Management",
        "Wireframing", "Communication", "SQL", "Market Research"
    ],
    "HR Executive": [
        "Recruitment", "Onboarding", "Employee Engagement", "HRIS",
        "Payroll", "Labor Law", "Performance Management", "Training",
        "Communication", "Conflict Resolution", "Excel", "LinkedIn"
    ],
    "Finance Analyst": [
        "Financial Modeling", "Excel", "SQL", "Accounting", "Budgeting",
        "Forecasting", "SAP", "Tableau", "Power BI", "Financial Reporting",
        "Risk Analysis", "VBA"
    ],
})

# ═══════════════════════════════════════════════════════════════
#  LANGUAGES LIST
# ═══════════════════════════════════════════════════════════════

LANGUAGES = [
    "English", "Hindi", "French", "German", "Spanish", "Japanese",
    "Chinese", "Korean", "Sanskrit", "Marathi", "Gujarati", "Tamil",
    "Telugu", "Bengali", "Punjabi", "Urdu", "Kannada", "Malayalam",
    "Odia", "Rajasthani",
]

# ═══════════════════════════════════════════════════════════════
#  COMMON MISSPELLINGS DICTIONARY
# ═══════════════════════════════════════════════════════════════

COMMON_MISSPELLINGS = {
    "pyhton": "Python",
    "pythn": "Python",
    "pytho": "Python",
    "phyton": "Python",
    "javscript": "JavaScript",
    "javascrpt": "JavaScript",
    "javasript": "JavaScript",
    "javacript": "JavaScript",
    "managment": "Management",
    "managemnt": "Management",
    "manegment": "Management",
    "comunication": "Communication",
    "comunicaton": "Communication",
    "communcation": "Communication",
    "communiction": "Communication",
    "experiance": "Experience",
    "exprience": "Experience",
    "experince": "Experience",
    "expereince": "Experience",
    "collage": "College",
    "colege": "College",
    "collge": "College",
    "acheivement": "Achievement",
    "acheivment": "Achievement",
    "achievment": "Achievement",
    "achivement": "Achievement",
    "certificaton": "Certification",
    "certifcation": "Certification",
    "certifiation": "Certification",
    "devlopment": "Development",
    "developement": "Development",
    "devloper": "Developer",
    "developr": "Developer",
    "programing": "Programming",
    "programmin": "Programming",
    "progamming": "Programming",
    "databse": "Database",
    "datbase": "Database",
    "enginnering": "Engineering",
    "enginering": "Engineering",
    "enginneering": "Engineering",
    "anaylsis": "Analysis",
    "analsis": "Analysis",
    "anlysis": "Analysis",
    "langugage": "Language",
    "langauge": "Language",
    "languag": "Language",
    "techonlogy": "Technology",
    "technlogy": "Technology",
    "tecnology": "Technology",
    "reume": "Resume",
    "resme": "Resume",
    "resumee": "Resume",
    "edcation": "Education",
    "educaton": "Education",
    "eduction": "Education",
    "prject": "Project",
    "porject": "Project",
    "projct": "Project",
    "respnsible": "Responsible",
    "responsble": "Responsible",
    "responisble": "Responsible",
    "implmented": "Implemented",
    "impelemented": "Implemented",
    "implemnted": "Implemented",
    "framwork": "Framework",
    "frmaework": "Framework",
    "fraemwork": "Framework",
    "oprating": "Operating",
    "opreating": "Operating",
    "operting": "Operating",
    "infromation": "Information",
    "informaton": "Information",
    "informatoin": "Information",
    "organzation": "Organization",
    "organisaton": "Organization",
    "orgnization": "Organization",
    "universtiy": "University",
    "univeristy": "University",
    "unversity": "University",
    "refrences": "References",
    "referances": "References",
    "refernces": "References",
    "intrested": "Interested",
    "intersted": "Interested",
    "intested": "Interested",
    "recieved": "Received",
    "recevied": "Received",
    "recived": "Received",
    "enviroment": "Environment",
    "environmnt": "Environment",
    "enviornment": "Environment",
    "strenght": "Strength",
    "stregth": "Strength",
    "strenth": "Strength",
}

# ═══════════════════════════════════════════════════════════════
#  CORRECT CAPITALIZATION MAP (technical terms)
# ═══════════════════════════════════════════════════════════════

CORRECT_CAPITALIZATION = {
    "python": "Python", "javascript": "JavaScript", "typescript": "TypeScript",
    "react": "React", "reactjs": "React.js", "react.js": "React.js",
    "angular": "Angular", "vue": "Vue", "vuejs": "Vue.js",
    "nodejs": "Node.js", "node.js": "Node.js",
    "html": "HTML", "css": "CSS", "sql": "SQL", "nosql": "NoSQL",
    "mongodb": "MongoDB", "mysql": "MySQL", "postgresql": "PostgreSQL",
    "github": "GitHub", "gitlab": "GitLab", "linkedin": "LinkedIn",
    "docker": "Docker", "kubernetes": "Kubernetes",
    "aws": "AWS", "azure": "Azure", "gcp": "GCP",
    "tensorflow": "TensorFlow", "pytorch": "PyTorch",
    "flask": "Flask", "django": "Django", "fastapi": "FastAPI",
    "java": "Java", "kotlin": "Kotlin", "swift": "Swift",
    "php": "PHP", "ruby": "Ruby", "golang": "GoLang",
    "jenkins": "Jenkins", "jira": "JIRA", "figma": "Figma",
    "excel": "Excel", "powerbi": "Power BI", "tableau": "Tableau",
    "selenium": "Selenium", "postman": "Postman",
    "android": "Android", "ios": "iOS", "linux": "Linux",
    "api": "API", "rest": "REST", "graphql": "GraphQL",
    "ci/cd": "CI/CD", "devops": "DevOps", "mlops": "MLOps",
    "opencv": "OpenCV", "pandas": "Pandas", "numpy": "NumPy",
    "scikit-learn": "Scikit-learn", "keras": "Keras",
    "matlab": "MATLAB", "sas": "SAS",
}


# ═══════════════════════════════════════════════════════════════
#  TEXT EXTRACTION
# ═══════════════════════════════════════════════════════════════

def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    import pdfplumber
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    # Clean up unmapped CID characters (like (cid:127)) which are usually bullets
                    page_text = re.sub(r"\(cid:\d+\)", "•", page_text)
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
    return text


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document
    text = ""
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
    return text


def extract_text(filepath: str) -> str:
    """Detect file type and extract text accordingly."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    return ""


# ═══════════════════════════════════════════════════════════════
#  FIELD EXTRACTION  (regex-based)
# ═══════════════════════════════════════════════════════════════

def extract_email(text: str) -> str:
    """Extract first email address from text."""
    match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    """Extract first phone number from text."""
    match = re.search(
        r"(?:\+?\d{1,3}[\s\-]?)?\(?\d{3,5}\)?[\s\-]?\d{3,4}[\s\-]?\d{3,4}", text
    )
    return match.group(0).strip() if match else ""


def extract_name(text: str) -> str:
    """Heuristic: first non-empty line is likely the candidate name."""
    for line in text.split("\n"):
        line = line.strip()
        # Skip lines that look like contact info or section headers
        if line and not re.match(r"^(email|phone|address|objective|summary|profile)", line, re.I):
            if not re.search(r"@|http|www\.|\.com|\.in|\.org", line, re.I):
                if len(line.split()) <= 5:
                    return line
    return ""


def extract_linkedin(text: str) -> str:
    """Extract LinkedIn profile URL."""
    match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9\-_/]+", text, re.I)
    return match.group(0) if match else ""


def extract_github(text: str) -> str:
    """Extract GitHub profile URL."""
    match = re.search(r"(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9\-_/]+", text, re.I)
    return match.group(0) if match else ""


# ═══════════════════════════════════════════════════════════════
#  SECTION DETECTION
# ═══════════════════════════════════════════════════════════════

def _section_present(text: str, keywords: list) -> bool:
    """Check if a section is present in the text by looking for header keywords."""
    text_lower = text.lower()
    for kw in keywords:
        # Look for keyword as a section heading (standalone on a line or followed by colon)
        pattern = rf"(?:^|\n)\s*{re.escape(kw)}\s*[:\-]?\s*(?:\n|$)"
        if re.search(pattern, text_lower):
            return True
    return False


def has_education(text: str) -> bool:
    return _section_present(text, ["education", "academic", "academics", "qualification", "qualifications", "degree"])


def has_experience(text: str) -> bool:
    return _section_present(text, ["experience", "work experience", "professional experience", "employment", "work history"])


def has_projects(text: str) -> bool:
    return _section_present(text, ["project", "projects", "academic projects", "personal projects", "key projects"])


def has_certifications(text: str) -> bool:
    return _section_present(text, ["certification", "certifications", "certificates", "certificate", "professional certifications"])


# ═══════════════════════════════════════════════════════════════
#  SKILL DETECTION
# ═══════════════════════════════════════════════════════════════

def detect_skills(text: str, role: str) -> dict:
    """
    Detect skills in text and compare against required skills for the given role.
    Returns dict with detected, required, matched, missing skills and match percentage.
    """
    required = JOB_ROLES.get(role, [])
    text_lower = text.lower()

    # Build a set of all known skills across all roles for general detection
    all_skills = set()
    for skills in JOB_ROLES.values():
        all_skills.update(skills)

    detected = []
    for skill in all_skills:
        # Use word boundary matching for short skill names
        pattern = rf"\b{re.escape(skill.lower())}\b"
        if re.search(pattern, text_lower):
            detected.append(skill)

    matched = [s for s in required if s in detected]
    missing = [s for s in required if s not in detected]

    match_pct = round((len(matched) / len(required)) * 100, 1) if required else 0

    return {
        "detected": sorted(detected),
        "required": required,
        "matched": sorted(matched),
        "missing": sorted(missing),
        "match_percentage": match_pct,
    }


# ═══════════════════════════════════════════════════════════════
#  LANGUAGE DETECTION
# ═══════════════════════════════════════════════════════════════

def detect_languages(text: str) -> list:
    """Detect human languages mentioned in the resume text."""
    found = []
    text_lower = text.lower()
    for lang in LANGUAGES:
        if re.search(rf"\b{re.escape(lang.lower())}\b", text_lower):
            found.append(lang)
    return found


# ═══════════════════════════════════════════════════════════════
#  FORMATTING ANALYSIS
# ═══════════════════════════════════════════════════════════════

def analyze_formatting(text: str, filepath: str) -> list:
    """
    Analyze resume formatting and return a list of suggestion strings.
    For DOCX files, also checks font sizes.
    """
    suggestions = []

    # ── Extra spaces ──
    if re.search(r"[a-zA-Z]  +[a-zA-Z]", text):
        suggestions.append("Remove extra spaces between words in your resume.")

    # ── Missing sections ──
    if not has_education(text):
        suggestions.append("Add a separate 'Education' section to your resume.")
    if not has_projects(text):
        suggestions.append("Add a 'Projects' section to showcase your work.")
    if not has_certifications(text):
        suggestions.append("Consider adding a 'Certifications' section if you have any.")
    if not has_experience(text):
        suggestions.append("Add a 'Work Experience' section to highlight your professional background.")

    # ── Word count check ──
    word_count = len(text.split())
    if word_count < 150:
        suggestions.append(f"Resume appears too short ({word_count} words). Aim for at least 300 words.")
    elif word_count > 1200:
        suggestions.append(f"Resume appears too long ({word_count} words). Try to keep it within 1–2 pages (300–1000 words).")

    # ── Capitalization of technical terms ──
    words_in_text = re.findall(r"\b[a-zA-Z.+#/\-]+\b", text)
    cap_issues = set()
    for word in words_in_text:
        lower = word.lower()
        if lower in CORRECT_CAPITALIZATION:
            correct = CORRECT_CAPITALIZATION[lower]
            if word != correct and word.lower() != word.upper():
                cap_issues.add(f"Use correct capitalization: '{correct}' instead of '{word}'.")
    suggestions.extend(sorted(cap_issues))

    # ── Contact info check ──
    if not extract_email(text):
        suggestions.append("Add your email address to the resume.")
    if not extract_phone(text):
        suggestions.append("Add your phone number to the resume.")
    if not extract_linkedin(text):
        suggestions.append("Add your LinkedIn profile link to the resume.")
    if not extract_github(text):
        suggestions.append("Consider adding your GitHub profile link.")

    # ── DOCX-specific font size checks ──
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".docx":
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document(filepath)
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.size:
                        pt = run.font.size.pt
                        # Detect heading-like text (short, bold, or large)
                        is_heading = run.bold or (len(para.text.strip().split()) <= 5 and pt >= 13)
                        if is_heading:
                            if pt < 14:
                                suggestions.append(f"Heading font size ({pt}pt) is too small. Use 14–16pt for headings.")
                                break
                            elif pt > 16:
                                suggestions.append(f"Heading font size ({pt}pt) is too large. Use 14–16pt for headings.")
                                break
                        else:
                            if pt < 10:
                                suggestions.append(f"Body font size ({pt}pt) is too small. Use 10–12pt for body text.")
                                break
                            elif pt > 12:
                                suggestions.append(f"Body font size ({pt}pt) is too large. Use 10–12pt for body text.")
                                break
        except Exception:
            pass

    # Deduplicate
    seen = set()
    unique = []
    for s in suggestions:
        if s not in seen:
            seen.add(s)
            unique.append(s)

    return unique


# ═══════════════════════════════════════════════════════════════
#  SPELLING CHECKER (dictionary-based fallback)
# ═══════════════════════════════════════════════════════════════

def check_spelling(text: str) -> list:
    """
    Check text against common misspelling dictionary.
    Returns list of suggestion strings.
    """
    suggestions = []
    words = re.findall(r"\b[a-zA-Z]+\b", text)
    found_mistakes = set()

    for word in words:
        lower = word.lower()
        if lower in COMMON_MISSPELLINGS and lower not in found_mistakes:
            correct = COMMON_MISSPELLINGS[lower]
            suggestions.append(f"'{word}' → '{correct}'")
            found_mistakes.add(lower)

    return suggestions


# ═══════════════════════════════════════════════════════════════
#  ATS SCORE CALCULATION
# ═══════════════════════════════════════════════════════════════

def calculate_ats_score(text: str, filepath: str, skill_result: dict) -> dict:
    """
    Calculate ATS score out of 100 with section-wise breakdown.

    Breakdown:
      Skills:         35 pts
      Education:      10 pts
      Projects:       15 pts
      Experience:     10 pts
      Certifications:  5 pts
      Contact:        10 pts
      Links:           5 pts
      Formatting:     10 pts
    ──────────────────────
      Total:         100 pts
    """
    breakdown = {}

    # ── Skills (max 35) ──
    match_pct = skill_result["match_percentage"]
    skills_score = round((match_pct / 100) * 35)
    breakdown["Skills"] = {"score": skills_score, "max": 35}

    # ── Education (max 10) ──
    edu_score = 10 if has_education(text) else 0
    breakdown["Education"] = {"score": edu_score, "max": 10}

    # ── Projects (max 15) ──
    proj_score = 15 if has_projects(text) else 0
    breakdown["Projects"] = {"score": proj_score, "max": 15}

    # ── Experience (max 10) ──
    exp_score = 10 if has_experience(text) else 0
    breakdown["Experience"] = {"score": exp_score, "max": 10}

    # ── Certifications (max 5) ──
    cert_score = 5 if has_certifications(text) else 0
    breakdown["Certifications"] = {"score": cert_score, "max": 5}

    # ── Contact details (max 10) ──
    contact_score = 0
    if extract_email(text):
        contact_score += 5
    if extract_phone(text):
        contact_score += 5
    breakdown["Contact"] = {"score": contact_score, "max": 10}

    # ── Links (max 5) ──
    links_score = 0
    if extract_linkedin(text):
        links_score += 3
    if extract_github(text):
        links_score += 2
    breakdown["Links"] = {"score": links_score, "max": 5}

    # ── Formatting quality (max 10) ──
    formatting_issues = analyze_formatting(text, filepath)
    # Start with full score and deduct per issue (max deduction = 10)
    fmt_deduction = min(len(formatting_issues) * 2, 10)
    fmt_score = max(10 - fmt_deduction, 0)
    breakdown["Formatting"] = {"score": fmt_score, "max": 10}

    total = sum(v["score"] for v in breakdown.values())
    total_max = sum(v["max"] for v in breakdown.values())

    return {
        "total_score": total,
        "total_max": total_max,
        "breakdown": breakdown,
    }


# ═══════════════════════════════════════════════════════════════
#  FULL ANALYSIS PIPELINE
# ═══════════════════════════════════════════════════════════════

def analyze_resume(filepath: str, selected_role: str) -> dict:
    """
    Run the full analysis pipeline on a resume file.
    Returns a dict with all extracted info, scores, and suggestions.
    """
    text = extract_text(filepath)
    if not text.strip():
        return {"error": "Could not extract text from the uploaded file."}

    # Field extraction
    candidate_name = extract_name(text)
    candidate_email = extract_email(text)
    candidate_phone = extract_phone(text)
    linkedin = extract_linkedin(text)
    github = extract_github(text)

    # Section detection
    edu = has_education(text)
    exp = has_experience(text)
    proj = has_projects(text)
    cert = has_certifications(text)

    # Skills
    skill_result = detect_skills(text, selected_role)

    # ATS Score
    ats_result = calculate_ats_score(text, filepath, skill_result)

    # Languages
    languages = detect_languages(text)

    # Formatting
    formatting_suggestions = analyze_formatting(text, filepath)

    # Spelling
    spelling_suggestions = check_spelling(text)

    # Final improvement suggestions
    final_suggestions = _generate_final_suggestions(
        skill_result, edu, proj, exp, cert, formatting_suggestions, spelling_suggestions
    )

    return {
        "candidate_name": candidate_name,
        "candidate_email": candidate_email,
        "candidate_phone": candidate_phone,
        "linkedin": linkedin,
        "github": github,
        "selected_role": selected_role,
        "education_found": edu,
        "experience_found": exp,
        "projects_found": proj,
        "certifications_found": cert,
        "detected_skills": skill_result["detected"],
        "matched_skills": skill_result["matched"],
        "missing_skills": skill_result["missing"],
        "required_skills": skill_result["required"],
        "job_match_percentage": skill_result["match_percentage"],
        "ats_score": ats_result["total_score"],
        "score_breakdown": ats_result["breakdown"],
        "languages_known": languages,
        "formatting_suggestions": formatting_suggestions,
        "spelling_suggestions": spelling_suggestions,
        "final_suggestions": final_suggestions,
        "raw_text": text,
    }


def _generate_final_suggestions(skill_result, edu, proj, exp, cert, fmt, spell) -> list:
    """Generate high-level improvement tips."""
    tips = []

    missing_count = len(skill_result["missing"])
    if missing_count > 0:
        tips.append(
            f"You are missing {missing_count} required skill(s) for this role. "
            f"Consider learning: {', '.join(skill_result['missing'][:5])}."
        )

    if skill_result["match_percentage"] < 50:
        tips.append("Your skill match is below 50%. Consider tailoring your resume for this specific role.")
    elif skill_result["match_percentage"] < 75:
        tips.append("Good skill match but there's room for improvement. Add more relevant skills.")

    if not edu:
        tips.append("An Education section is essential for most job applications.")
    if not proj:
        tips.append("Adding projects demonstrates practical experience and problem-solving ability.")
    if not exp:
        tips.append("Work experience is highly valued. Include internships or freelance work if applicable.")
    if not cert:
        tips.append("Relevant certifications can significantly boost your resume.")

    if len(fmt) > 3:
        tips.append("Multiple formatting issues detected. Review and clean up your resume layout.")

    if len(spell) > 0:
        tips.append("Fix spelling mistakes to maintain a professional impression.")

    if not tips:
        tips.append("Great resume! Keep it updated and tailored for each application.")

    return tips
