"""
resume_builder.py
-----------------
Generates a ready-to-share, ATS-friendly DOCX resume using the user's original text.
"""

import os
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════
#  HELPERS
# ═══════════════════════════════════════════════════════════════

def apply_spelling_corrections(text: str, suggestions: list) -> str:
    """Correct spelling mistakes in the raw text based on analysis suggestions."""
    if not text or not suggestions:
        return text
    for sug in suggestions:
        match = re.search(r"Misspelled '(.+?)'\. Did you mean '(.+?)'\?", sug)
        if match:
            wrong, right = match.groups()
            text = re.sub(r'\b' + re.escape(wrong) + r'\b', right, text, flags=re.IGNORECASE)
    return text


def extract_sections(text: str) -> dict:
    """Heuristically extract sections from the resume text."""
    # Clean up unmapped CID characters (like (cid:127)) which are usually bullets
    text = re.sub(r"\(cid:\d+\)", "•", text)

    sections = {
        "SUMMARY": "",
        "EXPERIENCE": "",
        "PROJECTS": "",
        "EDUCATION": "",
        "CERTIFICATIONS": "",
        "SKILLS": "",
        "LANGUAGES": "",
        "HOBBIES": "",
        "ACHIEVEMENTS": "",
        "CO_CURRICULAR": "",
        "PERSONAL_DETAILS": "",
        "UNKNOWN": ""
    }
    current_section = "UNKNOWN"
    
    headers = {
        r"^\s*(PROFESSIONAL\s+|WORK\s+)?EXPERIENCE\b": "EXPERIENCE",
        r"^\s*EMPLOYMENT(\s+HISTORY)?\b": "EXPERIENCE",
        r"^\s*WORK\s+HISTORY\b": "EXPERIENCE",
        
        r"^\s*(PERSONAL\s+|ACADEMIC\s+)?PROJECTS?\b": "PROJECTS",
        
        r"^\s*EDUCATION(AL)?(\s+QUALIFICATIONS?)?\b": "EDUCATION",
        r"^\s*ACADEMIC(\s+BACKGROUND|\s+QUALIFICATIONS?)?\b": "EDUCATION",
        
        r"^\s*CERTIFICATIONS?\b": "CERTIFICATIONS",
        r"^\s*COURSES?\b": "CERTIFICATIONS",
        r"^\s*TRAINING\b": "CERTIFICATIONS",
        
        r"^\s*(TECHNICAL\s+|IT\s+)?SKILLS?\b": "SKILLS",
        r"^\s*CORE\s+COMPETENCIES\b": "SKILLS",
        
        r"^\s*(PROFESSIONAL\s+|CAREER\s+)?SUMMARY\b": "SUMMARY",
        r"^\s*OBJECTIVE\b": "SUMMARY",
        r"^\s*PROFILE\b": "SUMMARY",
        
        r"^\s*LANGUAGES?\b": "LANGUAGES",
        r"^\s*HOBBIES(\s*&*\s*INTERESTS)?\b": "HOBBIES",
        r"^\s*INTERESTS?\b": "HOBBIES",
        r"^\s*ACHIEVEMENTS?\b": "ACHIEVEMENTS",
        r"^\s*(EXTRA\s*)?CURRICULAR(\s+ACTIVITIES)?\b": "ACHIEVEMENTS",
        r"^\s*CO[- ]?CURRICULAR(\s+ACTIVITIES)?\b": "CO_CURRICULAR",
        r"^\s*PERSONAL\s+(DETAILS|INFO|INFORMATION)\b": "PERSONAL_DETAILS"
    }
    
    for line in text.split('\n'):
        line_clean = line.strip().upper()
        # Clean non-alphanumeric at start of line for matching (e.g. bullets)
        line_match = re.sub(r'^[^A-Z]+', '', line_clean)
        
        matched = False
        for pattern, sec_name in headers.items():
            if re.match(pattern, line_match) and len(line_clean) < 40:
                current_section = sec_name
                matched = True
                break
        
        if not matched and line.strip():
            sections[current_section] += line.strip() + "\n"
            
    return sections


def parse_list(val):
    """Parse a value that may be a JSON string, a list, or a semicolon-separated string."""
    if not val:
        return []
    if isinstance(val, list):
        return val
    try:
        return json.loads(val)
    except Exception:
        return [v.strip() for v in val.split(";") if v.strip()]


# ═══════════════════════════════════════════════════════════════
#  DOCUMENT FORMATTING HELPERS
# ═══════════════════════════════════════════════════════════════

FONT_NAME = "Calibri"
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_DARK_GRAY = RGBColor(51, 51, 51)
COLOR_HEADING = RGBColor(0, 0, 0)


def set_paragraph_font(paragraph, size=10.5, bold=False, color=None):
    """Ensure every run in a paragraph uses the correct font."""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        run.font.color.rgb = color or COLOR_DARK_GRAY
        run.bold = bold
        # Force the font for East Asian text as well
        r_element = run._element
        rPr = r_element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)


def add_horizontal_line(doc):
    """Add a thin horizontal rule using a bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')        # thin line
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_section_heading(doc, title):
    """Add a bold uppercase section heading with a bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_HEADING
    # Force font
    r_element = run._element
    rPr = r_element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)

    # Add bottom border to the heading paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_body_paragraph(doc, text, bold=False, indent=False):
    """Add a properly formatted body paragraph with tight spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_DARK_GRAY
    run.bold = bold
    # Force font
    r_element = run._element
    rPr = r_element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    return p


def add_bullet_paragraph(doc, text):
    """Add a bullet point with tight spacing."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.4)
    # Clear default run and add our own
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_DARK_GRAY
    r_element = run._element
    rPr = r_element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    return p


def add_inline_label(doc, label, value, space_after=1):
    """Add a line like 'Label: value' with the label bolded."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(space_after)

    label_run = p.add_run(label)
    label_run.bold = True
    label_run.font.name = FONT_NAME
    label_run.font.size = Pt(10.5)
    label_run.font.color.rgb = COLOR_BLACK

    value_run = p.add_run(value)
    value_run.font.name = FONT_NAME
    value_run.font.size = Pt(10.5)
    value_run.font.color.rgb = COLOR_DARK_GRAY

    for run in p.runs:
        r_element = run._element
        rPr = r_element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)
    return p


# ═══════════════════════════════════════════════════════════════
#  SECTION CONTENT RENDERER
# ═══════════════════════════════════════════════════════════════

BULLET_CHARS = set('-*•·➢✓◦►▪')


def is_subheading(line: str) -> bool:
    """Determine if a line is likely a subheading (project name, job title, etc.)."""
    line = line.strip()
    if not line:
        return False
    # Must be short, not end with period, and not start with a bullet
    if len(line) > 70 or line.endswith('.'):
        return False
    if line[0] in BULLET_CHARS:
        return False
    # Looks like a title if it contains a dash separator (e.g. "MEDHA – Web App")
    # or is mostly title case / uppercase
    if '–' in line or '—' in line or ' - ' in line:
        return True
    # Check if most words are capitalized
    words = line.split()
    if len(words) <= 8:
        capitalized = sum(1 for w in words if w[0].isupper() or w in ('and', 'of', 'the', 'in', 'for', 'with', 'a', 'an', 'at', 'to', 'from'))
        if capitalized >= len(words) * 0.6:
            return True
    return False


def add_section_content(doc, content, section_name=""):
    """Render section content with proper formatting for subheadings, bullets, and body text."""
    content = content.strip()
    if not content:
        return False
    
    is_edu = section_name == "EDUCATION"
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Bullet points
        if line[0] in BULLET_CHARS:
            clean = line.lstrip(''.join(BULLET_CHARS) + ' ').strip()
            if clean:
                add_bullet_paragraph(doc, clean)
        # Subheadings (project names, job titles, etc.)
        elif is_subheading(line) and not is_edu:
            add_body_paragraph(doc, line, bold=True)
        # Regular body text
        else:
            # Auto-bullet long descriptive lines that aren't key-value pairs (like "CGPA: 8.9")
            is_kv = ':' in line and line.index(':') < 30
            if not is_edu and len(line.split()) >= 5 and not is_kv:
                add_bullet_paragraph(doc, line)
            else:
                add_body_paragraph(doc, line)
    
    return True


# ═══════════════════════════════════════════════════════════════
#  MAIN GENERATOR
# ═══════════════════════════════════════════════════════════════

def generate_corrected_resume(analysis_data: dict, output_dir: str, analysis_id: int) -> str:
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, f"corrected_resume_{analysis_id}.docx")
    doc = Document()
    
    # ── Page setup: standard letter margins ──
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ── Default style ──
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLOR_DARK_GRAY
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    # Force font in Normal style XML
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)

    # ── Set white background ──
    bg = doc.element.find(qn('w:background'))
    if bg is None:
        bg = OxmlElement('w:background')
        doc.element.insert(0, bg)
    bg.set(qn('w:color'), 'FFFFFF')

    # Parse analysis data
    raw_text = analysis_data.get("raw_text", "")
    spell_sugs = parse_list(analysis_data.get("spelling_suggestions"))
    corrected_text = apply_spelling_corrections(raw_text, spell_sugs)
    original_sections = extract_sections(corrected_text)

    # ════════════════════════════════════════════
    #  1. HEADER — Name & Contact
    # ════════════════════════════════════════════
    name = analysis_data.get("candidate_name", "Your Name")
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(name.upper() if name else "YOUR NAME")
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.name = FONT_NAME
    name_run.font.color.rgb = COLOR_BLACK

    # Contact info line
    contact_info = []
    email = analysis_data.get("candidate_email")
    phone = analysis_data.get("candidate_phone")
    linkedin = analysis_data.get("linkedin")
    github = analysis_data.get("github")
    if email:
        contact_info.append(email)
    if phone:
        contact_info.append(phone)
    if linkedin:
        contact_info.append(linkedin)
    if github:
        contact_info.append(github)

    if contact_info:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_before = Pt(0)
        contact_para.paragraph_format.space_after = Pt(2)
        contact_run = contact_para.add_run(" | ".join(contact_info))
        contact_run.font.name = FONT_NAME
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = COLOR_DARK_GRAY

    # Thin horizontal line under header
    add_horizontal_line(doc)

    # ════════════════════════════════════════════
    #  2. PROFESSIONAL SUMMARY
    # ════════════════════════════════════════════
    summary_text = original_sections["SUMMARY"].strip()
    if not summary_text:
        role = analysis_data.get('selected_role', 'Professional')
        summary_text = f"Results-driven {role} with experience leveraging technical skills to drive business outcomes."
    
    add_section_heading(doc, "Professional Summary")
    add_body_paragraph(doc, summary_text)

    # ════════════════════════════════════════════
    #  3. TECHNICAL SKILLS
    # ════════════════════════════════════════════
    matched = parse_list(analysis_data.get("matched_skills"))
    missing = parse_list(analysis_data.get("missing_skills"))
    detected = parse_list(analysis_data.get("detected_skills"))
    languages_known = parse_list(analysis_data.get("languages_known"))

    add_section_heading(doc, "Technical Skills")

    all_good_skills = matched if matched else detected
    if all_good_skills:
        add_inline_label(doc, "Core Competencies: ", ", ".join(all_good_skills))

    if missing:
        add_inline_label(doc, "Suggested Skills: ", ", ".join(missing))

    if languages_known or original_sections["LANGUAGES"].strip():
        combined_langs = set(languages_known)
        for lang in original_sections["LANGUAGES"].split('\n'):
            if lang.strip():
                combined_langs.add(lang.strip())
        add_inline_label(doc, "Languages: ", ", ".join(sorted(combined_langs)))

    # ════════════════════════════════════════════
    #  4. PROFESSIONAL EXPERIENCE
    # ════════════════════════════════════════════
    exp_content = original_sections["EXPERIENCE"].strip()
    if exp_content:
        add_section_heading(doc, "Professional Experience")
        add_section_content(doc, exp_content)
    elif analysis_data.get("experience_found"):
        unknown_content = original_sections["UNKNOWN"].strip()
        if unknown_content:
            add_section_heading(doc, "Professional Experience")
            add_section_content(doc, unknown_content)

    # ════════════════════════════════════════════
    #  5. PROJECTS
    # ════════════════════════════════════════════
    proj_content = original_sections["PROJECTS"].strip()
    if proj_content:
        add_section_heading(doc, "Projects")
        add_section_content(doc, proj_content)

    # ════════════════════════════════════════════
    #  6. EDUCATION
    # ════════════════════════════════════════════
    edu_content = original_sections["EDUCATION"].strip()
    if edu_content:
        add_section_heading(doc, "Education")
        add_section_content(doc, edu_content, section_name="EDUCATION")

    # ════════════════════════════════════════════
    #  7. CERTIFICATIONS
    # ════════════════════════════════════════════
    cert_content = original_sections["CERTIFICATIONS"].strip()
    if cert_content:
        add_section_heading(doc, "Certifications")
        add_section_content(doc, cert_content)

    # ════════════════════════════════════════════
    #  8. ACHIEVEMENTS
    # ════════════════════════════════════════════
    ach_content = original_sections["ACHIEVEMENTS"].strip()
    if ach_content:
        add_section_heading(doc, "Achievements")
        add_section_content(doc, ach_content)

    # ════════════════════════════════════════════
    #  9. HOBBIES & INTERESTS
    # ════════════════════════════════════════════
    hob_content = original_sections["HOBBIES"].strip()
    if hob_content:
        add_section_heading(doc, "Hobbies & Interests")
        add_section_content(doc, hob_content)

    # ════════════════════════════════════════════
    #  10. CO-CURRICULAR ACTIVITIES
    # ════════════════════════════════════════════
    co_content = original_sections.get("CO_CURRICULAR", "").strip()
    if co_content:
        add_section_heading(doc, "Co-Curricular Activities")
        add_section_content(doc, co_content)

    # ════════════════════════════════════════════
    #  11. PERSONAL DETAILS
    # ════════════════════════════════════════════
    personal_content = original_sections.get("PERSONAL_DETAILS", "").strip()
    if personal_content:
        add_section_heading(doc, "Personal Details")
        add_section_content(doc, personal_content)

    doc.save(filepath)
    return filepath
